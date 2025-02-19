from docx import Document
from docx.shared import Inches
import os
from pygui import *


def pDocx(ClientName, ClientID,current_date,response,serviceProvided,serviceProvidedBy,startTime,endTime,totaltime,timeQuarter):
    document = Document()

# header organization
    p = document.add_paragraph('Service Log' +  '\n' + current_date.strftime('%m-%d-%Y') + "\n")
    p.add_run(ClientName + "\n")
    p.add_run('Medicaid Number: ' + ClientID + "\n")
    p.add_run(serviceProvided + "\n")
    p.add_run(serviceProvidedBy + "\n")
    p.add_run('Peterson Family Care LLC Provider Number: 009279700'+ "\n")
    p.add_run('Start Time: ' + str(startTime) + ' - End Time: '+ str(endTime) + "\n")
    p.add_run('Total Hours: ' + str(totaltime) + "\n")
    p.add_run('Total Quarter Hours: ' + str(timeQuarter) + "\n")



# AI generated Body Paragraph
    document.add_heading('Daily Report', level=1)
    document.add_paragraph(response.message.content + "\n")


# end document signatures
    signature = document.add_paragraph('Client Signature: ')
    try:
        names = ClientName.split(' ')
        sig = ''
        for name in names:
            sig = sig + name[0]
        signature.add_run(sig + '\n').font.name='Segoe Script'
        

    except Exception as e:
        popup_err(e)

    signature.add_run('Provider Signature: ')
    signature.add_run(serviceProvidedBy).font.name = 'Segoe Script'
    


    if not os.path.exists(ClientName):
        os.mkdir(ClientName)
    
    document.save(str(ClientName) + '/' + str(ClientName).replace(" ","-") +' ' + current_date.strftime('%m-%d-%Y') + ".docx")